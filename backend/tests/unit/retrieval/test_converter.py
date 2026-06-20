import io

import pytest
from docx import Document as DocxDocument

from retrieval.converter import convert_to_markdown


def test_passthrough_for_markdown():
    assert convert_to_markdown("notes.md", b"# Title\n\nbody") == "# Title\n\nbody"


def test_unsupported_extension_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        convert_to_markdown("sheet.xlsx", b"whatever")


def test_docx_is_converted_to_markdown():
    doc = DocxDocument()
    doc.add_heading("Market Overview", level=1)
    doc.add_paragraph("Demand for mining equipment is rising.")
    buf = io.BytesIO()
    doc.save(buf)

    markdown = convert_to_markdown("report.docx", buf.getvalue())

    assert "Market Overview" in markdown
    assert "Demand for mining equipment is rising." in markdown
